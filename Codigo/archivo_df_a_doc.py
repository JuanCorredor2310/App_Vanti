from docx import Document
from docx.shared import Inches, Pt
from docx2pdf import convert
import os
from contextlib import redirect_stdout, redirect_stderr
from win32com.client import Dispatch
import pythoncom

def add_horizontal_line(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.text = "\n"
    paragraph.add_run().add_break()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.border_bottom = True
    run.text = "-"*2850
    run.font.size = Pt(1)

def cerrar_documentos_word(thread=None):
    try:
        pythoncom.CoInitialize()  # Inicializa el modelo COM
        word = Dispatch("Word.Application")
        if word.Documents.Count > 0:
            for i in range(word.Documents.Count, 0, -1):
                doc = word.Documents(i)
                doc.Close(SaveChanges=False)
        word.Quit()
    except BaseException as e:
        if thread:
            thread.message_sent.emit(f"Error al cerrar documentos de Word: {str(e)}", "red")
        else:
            print(f"Error al cerrar documentos de Word: {str(e)}")
    finally:
        pythoncom.CoUninitialize()  # Libera el modelo COM

def almacenar_errores(dic_errores, filial, c_filial, mes, anio, nombre, largo, num, clasi, thread=None):
    try:
        pythoncom.CoInitialize()  # Inicializa el modelo COM al inicio de la función

        cerrar_documentos_word(thread=thread)  # Cierra cualquier documento abierto de Word

        doc = Document()
        section = doc.sections[0]
        section.orientation = 1
        section.page_height = Inches(largo * 4.5)
        section.page_width = Inches(16)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        parrafo = doc.add_paragraph()
        parrafo.add_run(f"Errores encontrados para el reporte {num} de Desviaciones Significativas {clasi}").bold = True
        parrafo = doc.add_paragraph()
        parrafo.add_run(f"{filial} ({c_filial}) para {mes}-{anio}").bold = True

        add_horizontal_line(doc)

        for llave, valor in dic_errores.items():
            df = valor[1].copy()
            columna = len(df.columns)
            parrafo = doc.add_paragraph()
            parrafo.add_run(valor[0]).bold = True

            table = doc.add_table(rows=1, cols=len(df.columns))
            hdr_cells = table.rows[0].cells
            for i, column_name in enumerate(df.columns):
                hdr_cells[i].text = column_name
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        if i == columna - 1:
                            run.bold = True

            for row in df.itertuples(index=False):
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    paragraph = row_cells[i].paragraphs[0]
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = Pt(1)
                    run = row_cells[i].add_paragraph().add_run(str(value))
                    run.font.size = Pt(11)
                    if i == columna - 1:
                        run.bold = True

            add_horizontal_line(doc)

        try:
            doc.save(nombre)
            with open(os.devnull, 'w') as fnull:
                with redirect_stdout(fnull), redirect_stderr(fnull):
                    convert(nombre)
            if os.path.exists(nombre):
                os.remove(nombre)
            return True
        except BaseException as e:
            if thread:
                thread.message_sent.emit(f"Error al generar el reporte: {str(e)}", "red")
            else:
                print(f"Error al generar el reporte: {str(e)}")
            return False

    except Exception as e:
        if thread:
            thread.message_sent.emit(f"Error general: {str(e)}", "red")
        else:
            print(f"Error general: {str(e)}")
        return False
    finally:
        pythoncom.CoUninitialize()  # Libera el modelo COM al finalizar

